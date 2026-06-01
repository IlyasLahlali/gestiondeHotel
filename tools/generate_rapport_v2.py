# -*- coding: utf-8 -*-
"""Génère le rapport de projet HôtelFacile Smart fusionné (v2) dans docs/Rapport.docx"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r"C:\Users\ilyas\Desktop\gestion-hotel\docs\Rapport.docx"

BLUE = RGBColor(0x2F, 0x64, 0xEA)
GOLD = RGBColor(0xFB, 0xBF, 0x24)
DARK = RGBColor(0x11, 0x18, 0x27)


def set_cell_shading(cell, fill_hex: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_placeholder(doc, title, hint, height_lines=2):
    p = doc.add_paragraph()
    run = p.add_run(f"【 {title} 】")
    run.bold = True
    run.font.color.rgb = BLUE
    for _ in range(height_lines):
        box = doc.add_paragraph()
        box.add_run(hint).italic = True
        box.paragraph_format.left_indent = Cm(1)
        box.paragraph_format.space_after = Pt(6)
    doc.add_paragraph()


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = BLUE if level == 1 else DARK
    return h


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Cm(1 + level * 0.5)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_table(doc, headers, rows, header_fill="EFF6FF"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], header_fill)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph()
    return table


def add_actor(doc, title, desc):
    p = doc.add_paragraph()
    p.add_run(f"{title} — ").bold = True
    p.add_run(desc)
    p.paragraph_format.space_after = Pt(6)


def add_user_story(doc, code, role, want, goal):
    add_heading(doc, code, 3)
    add_body(doc, f"En tant que {role},")
    add_body(doc, f"Je veux {want},")
    add_body(doc, f"Afin de {goal}.")


def build():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2.5)
        s.bottom_margin = Cm(2.5)
        s.left_margin = Cm(2.5)
        s.right_margin = Cm(2.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # ===================== PAGE DE GARDE =====================
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Université Privée de Marrakech (UPM)")
    r.font.size = Pt(12)
    r.font.color.rgb = DARK

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Filière : Développement Logiciel & Sécurité Informatique").font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Année universitaire : 2025 – 2026").font.size = Pt(11)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("RAPPORT DE PROJET DE FIN D'ÉTUDES")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = DARK

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("HôtelFacile Smart")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = BLUE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = p.add_run("Plateforme web intelligente de gestion et de réservation d'hôtels au Maroc")
    rs.font.size = Pt(14)
    rs.font.color.rgb = DARK

    doc.add_paragraph()
    cover = [
        ("Spécialité", "Développement Logiciel & Sécurité Informatique"),
        ("Thématique", "Tourisme / Hôtellerie numérique"),
        ("Réalisé par", "[Nom Prénom — à compléter]\n[Nom Prénom — à compléter]"),
        ("Encadré par", "[Nom de l'enseignant — à compléter]"),
        ("Établissement", "Université Privée de Marrakech (UPM)"),
        ("Année universitaire", "2025 – 2026"),
    ]
    for label, value in cover:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"{label} : ").bold = True
        p.add_run(value)

    add_placeholder(doc, "LOGO ÉTABLISSEMENT / LOGO PROJET",
                    "Insérer le logo UPM et/ou le logo HôtelFacile Smart (frontend/images/logo.svg).", 1)
    doc.add_page_break()

    # ===================== REMERCIEMENTS =====================
    add_heading(doc, "Remerciements", 1)
    add_body(doc,
        "Nous tenons en premier lieu à exprimer notre profonde gratitude envers notre encadrant, "
        "dont la disponibilité, la rigueur et les précieux conseils ont guidé notre travail tout "
        "au long de ce projet. Son accompagnement bienveillant a été une source constante de "
        "motivation et d'inspiration.")
    add_body(doc,
        "Nous remercions également l'ensemble du corps enseignant de l'Université Privée de "
        "Marrakech (UPM), filière Développement Logiciel & Sécurité Informatique, pour la "
        "formation de qualité qu'il nous a dispensée et pour les valeurs académiques transmises.")
    add_body(doc,
        "Nos remerciements vont aussi à nos familles et à nos proches pour leur soutien "
        "indéfectible, leur patience et leurs encouragements durant cette période intense de "
        "conception, de codage et de tests.")
    add_body(doc,
        "Enfin, nous remercions tous ceux qui, de près ou de loin, ont contribué à la "
        "réalisation de HôtelFacile Smart, que ce soit par leurs retours, leurs idées ou "
        "leur aide technique.")
    doc.add_page_break()

    # ===================== RÉSUMÉ =====================
    add_heading(doc, "Résumé (Abstract)", 1)

    add_heading(doc, "Résumé — Français", 2)
    add_body(doc,
        "Ce projet de fin d'études consiste à concevoir et développer HôtelFacile Smart, "
        "une application web full-stack dédiée à la gestion et à la réservation d'hôtels au Maroc. "
        "La plateforme met en relation quatre profils — visiteur, client voyageur, propriétaire "
        "d'établissement et administrateur — au sein d'un écosystème unique, sécurisé et moderne.")
    add_body(doc,
        "La problématique traitée est la fragmentation des outils utilisés par les hôteliers "
        "indépendants et les voyageurs : réservation manuelle, absence de visibilité en ligne, "
        "modération des établissements, suivi des chambres et des avis clients. La solution "
        "proposée centralise l'inscription, la recherche multi-critères, la réservation, la "
        "modération administrative, les notifications en temps réel et la gestion des favoris et avis.")
    add_body(doc,
        "Les technologies mobilisées comprennent Node.js et Express.js pour le backend REST, "
        "MySQL pour la persistance, HTML/CSS/JavaScript vanilla pour le frontend multi-portail, "
        "JWT et bcrypt pour la sécurité, Leaflet/OpenStreetMap pour la géolocalisation, et "
        "Google OAuth 2.0 pour l'authentification sociale.")
    add_body(doc,
        "Les résultats obtenus montrent qu'il est possible de livrer une application cohérente, "
        "responsive et évolutive, avec une charte graphique unifiée (bleu #2F64EA et jaune "
        "#FBBF24), des parcours utilisateurs distincts et une architecture monolithique légère "
        "adaptée à un contexte académique.")

    add_heading(doc, "Abstract — English", 2)
    add_body(doc,
        "This final year project presents the design and development of HôtelFacile Smart, "
        "a full-stack web application dedicated to hotel management and booking in Morocco. "
        "The platform connects four user profiles — visitor, traveler client, hotel owner and "
        "administrator — within a single, secure and modern ecosystem.")
    add_body(doc,
        "The solution is built on a Node.js/Express.js backend, MySQL database, and a multi-portal "
        "vanilla HTML/CSS/JS frontend. Key features include JWT and Google OAuth 2.0 authentication, "
        "real-time notifications, hotel geolocation, multimedia galleries, review and rating systems, "
        "and statistical dashboards for owners and administrators.")

    add_heading(doc, "Mots-clés", 2)
    add_body(doc,
        "Gestion hôtelière • Application web • Node.js • Express.js • MySQL • JWT • OAuth 2.0 • "
        "Réservation en ligne • Maroc • Full-Stack • Tourisme numérique")
    doc.add_page_break()

    # ===================== TABLE DES MATIÈRES =====================
    add_heading(doc, "Table des matières", 1)
    add_body(doc,
        "Générer automatiquement dans Word : Références → Table des matières → Table automatique. "
        "Mettre à jour après insertion de toutes les figures et captures.")
    for item in [
        "Introduction générale",
        "Chapitre 1 — Contexte et état de l'art",
        "Chapitre 2 — Analyse des besoins",
        "Chapitre 3 — Conception",
        "Chapitre 4 — Développement",
        "Chapitre 5 — Résultats et perspectives",
        "Conclusion générale",
        "Bibliographie",
        "Annexes",
    ]:
        add_bullet(doc, item)
    doc.add_page_break()

    # ===================== INTRODUCTION =====================
    add_heading(doc, "Introduction générale", 1)

    add_heading(doc, "Contexte du projet", 2)
    add_body(doc,
        "Le secteur touristique marocain occupe une place stratégique dans l'économie nationale. "
        "Avec des millions de visiteurs chaque année et des destinations phares comme Marrakech, "
        "Fès, Casablanca, Rabat, Tanger, Agadir, Oujda et Meknès, la demande de solutions "
        "numériques pour découvrir, comparer et réserver des hébergements ne cesse de croître.")
    add_body(doc,
        "Malgré cet essor, une part significative des établissements hôteliers indépendants "
        "souffre encore d'un manque d'outils numériques adaptés : gestion des chambres, des "
        "réservations, de la visibilité en ligne et des performances. C'est dans ce contexte "
        "que s'inscrit HôtelFacile Smart, une plateforme visant à combler ce fossé digital.")

    add_heading(doc, "Présentation du domaine", 2)
    add_body(doc,
        "L'e-tourisme désigne l'ensemble des usages du numérique appliqués au secteur touristique. "
        "Des plateformes internationales comme Booking.com, Airbnb ou Expedia ont transformé les "
        "habitudes de voyage. Au Maroc, la transformation numérique est en cours mais inégale : "
        "si les grandes enseignes disposent de PMS (Property Management Systems) sophistiqués, "
        "les hôteliers indépendants n'ont pas toujours les moyens de s'équiper.")
    add_body(doc,
        "Notre application se positionne à l'intersection des OTA (Online Travel Agencies) et "
        "des outils de gestion interne, en proposant une plateforme unifiée couvrant la vitrine "
        "publique, l'espace client, l'espace propriétaire et le back-office administrateur.")

    add_heading(doc, "Problématique", 2)
    add_body(doc,
        "Comment permettre à un hôtelier indépendant au Maroc de gérer efficacement son "
        "établissement en ligne — de la publication de ses chambres à la gestion des réservations "
        "— sans recourir à des solutions coûteuses ou complexes ?")
    add_body(doc,
        "Et du côté du voyageur : comment lui offrir une expérience de recherche, de comparaison "
        "et de réservation fluide, fiable et sécurisée, centrée sur les destinations marocaines, "
        "tout en garantissant aux administrateurs un contrôle centralisé de la qualité des offres ?")

    add_heading(doc, "Objectifs", 2)
    for o in [
        "Développer une application web multi-rôles (Public, Client, Propriétaire, Admin).",
        "Implémenter une API REST complète avec authentification JWT et contrôle d'accès par rôle.",
        "Concevoir une base de données relationnelle MySQL normalisée et évolutive.",
        "Proposer une interface moderne, cohérente et responsive (charte #2F64EA / #FBBF24).",
        "Intégrer des fonctionnalités « intelligentes » : recherche avancée, hôtels populaires, "
        "recherche floue admin, fusion recherche propriétaire, notifications contextuelles.",
        "Documenter l'architecture, les choix techniques et les scénarios de test.",
    ]:
        add_bullet(doc, o)

    add_heading(doc, "Structure du rapport", 2)
    add_body(doc,
        "Ce document est organisé en cinq chapitres conformément au canevas pédagogique. "
        "Le chapitre 1 présente le contexte et l'état de l'art. Le chapitre 2 détaille "
        "l'analyse des besoins. Le chapitre 3 expose la conception. Le chapitre 4 décrit "
        "le développement et la sécurité. Le chapitre 5 analyse les résultats, les tests "
        "et les perspectives d'évolution.")
    doc.add_page_break()

    # ===================== CHAPITRE 1 =====================
    add_heading(doc, "Chapitre 1 : Contexte et état de l'art", 1)

    add_heading(doc, "1.1 Présentation du domaine", 2)
    add_body(doc,
        "Le secteur hôtelier regroupe l'ensemble des établissements d'hébergement payant. "
        "Sa gestion implique des processus interdépendants : disponibilités, réservations, "
        "facturation, relation client et maintenance. La digitalisation a conduit à deux "
        "grandes catégories d'outils :")
    add_bullet(doc, "Les PMS : logiciels de gestion interne pour les hôtels (réception, facturation).")
    add_bullet(doc, "Les OTA : agences de voyage en ligne permettant la distribution et la réservation.")
    add_body(doc,
        "Au Maroc, la diversité des types d'hébergement (économique, standard, supérieure, "
        "deluxe, suite, familiale) et la répartition géographique des villes touristiques "
        "imposent des filtres performants et une cartographie claire. HôtelFacile Smart "
        "cible explicitement ce contexte avec huit villes marocaines supportées.")

    add_heading(doc, "1.2 Problématique", 2)
    add_heading(doc, "1.2.1 Description du problème", 3)
    add_body(doc,
        "Le principal problème identifié est la fracture numérique au sein du secteur hôtelier "
        "marocain. Les hôteliers indépendants gèrent encore leurs réservations par téléphone, "
        "e-mail ou tableurs Excel, sans visibilité structurée. Les voyageurs naviguent entre "
        "plusieurs sites internationaux qui ne couvrent pas toutes les offres locales.")
    add_heading(doc, "1.2.2 Impact", 3)
    add_bullet(doc, "Propriétaires : perte de clients, gestion chronophage, absence d'analytique.")
    add_bullet(doc, "Clients : recherche fragmentée, manque de transparence, risque de double réservation.")
    add_bullet(doc, "Secteur : ralentissement de la digitalisation du tourisme marocain.")
    add_heading(doc, "1.2.3 Importance", 3)
    add_body(doc,
        "Avec l'essor du tourisme au Maroc, une solution numérique adaptée aux PME hôtelières "
        "peut avoir un impact économique significatif sur les petites et moyennes structures.")

    add_heading(doc, "1.3 Étude des solutions existantes", 2)
    add_body(doc, "Analyse comparative des principales plateformes de référence :")
    add_table(doc,
        ["Solution", "Type", "Avantages", "Limites"],
        [
            ("Booking.com", "OTA internationale", "Large audience, UX soignée", "Commissions 15–25 %, peu adapté aux PME locales"),
            ("Airbnb", "OTA pair-à-pair", "Communauté forte, UX fluide", "Orienté hébergement chez l'habitant"),
            ("Expedia", "OTA internationale", "Forte visibilité", "Complexité pour hôteliers locaux"),
            ("Opera PMS", "PMS professionnel", "Très complet", "Coûteux, complexe, inadapté aux PME"),
            ("HôtelFacile Smart", "Plateforme locale", "Gratuit, multi-rôle, modération intégrée", "Périmètre académique, 8 villes"),
        ])

    add_heading(doc, "1.4 Solution proposée", 2)
    add_heading(doc, "1.4.1 Idée principale", 3)
    add_body(doc,
        "HôtelFacile Smart est une marketplace hôtelière full-stack permettant aux hôteliers "
        "marocains de publier et gérer leurs établissements, et aux clients de rechercher, "
        "comparer et réserver des chambres dans les principales villes du Maroc.")
    add_heading(doc, "1.4.2 Valeur ajoutée", 3)
    for v in [
        "Centrage marché marocain : 8 villes, interface en français.",
        "Gratuité : aucune commission sur les réservations.",
        "Multi-portail : quatre interfaces dédiées (Public, Client, Propriétaire, Admin).",
        "Modération admin avant publication des hôtels.",
        "Notifications en temps réel et géolocalisation Leaflet.",
        "Workflow complet de réservation (en attente → confirmée / refusée / annulée).",
    ]:
        add_bullet(doc, v)
    add_heading(doc, "1.4.3 Apport technologique et « Smart »", 3)
    add_body(doc,
        "Le terme Smart reflète des mécanismes heuristiques et métier plutôt qu'un module "
        "d'apprentissage automatique au sens strict : classement des hôtels populaires, "
        "scoring par avis clients, recherche floue admin (normalisation accents, filtres combinés), "
        "recherche fusionnée avec le bandeau d'accueil propriétaire, politique de suppression "
        "intelligente (deletionPolicy.js), vérification des chevauchements de dates, et "
        "notifications automatiques déclenchées par les événements métier.")
    add_body(doc,
        "Une évolution future pourrait intégrer un moteur de recommandation par apprentissage "
        "automatique (filtrage collaboratif, prédiction d'occupation, détection d'anomalies sur les avis).")

    add_placeholder(doc, "FIGURE 1.1 — Schéma du positionnement de HôtelFacile Smart",
                    "Diagramme de contexte : acteurs externes, plateforme, base de données, services OAuth/OSM.", 2)
    doc.add_page_break()

    # ===================== CHAPITRE 2 =====================
    add_heading(doc, "Chapitre 2 : Analyse des besoins", 1)

    add_heading(doc, "2.1 Identification des acteurs", 2)
    add_actor(doc, "Visiteur (Public)",
        "Consulte l'accueil, recherche des hôtels validés, explore les villes sans compte. "
        "Ne peut pas réserver ni accéder aux fonctionnalités personnalisées.")
    add_actor(doc, "Client",
        "Voyageur authentifié : recherche, réservation, favoris, avis, historique des séjours, "
        "notifications sur l'état de ses réservations.")
    add_actor(doc, "Propriétaire d'hôtel",
        "Hôtelier inscrit : ajoute et gère hôtels/chambres/photos, valide ou refuse les "
        "réservations, consulte avis et tableau de bord statistique (revenus, occupation).")
    add_actor(doc, "Administrateur",
        "Superutilisateur : valide ou refuse les hôtels soumis, consulte statistiques globales, "
        "effectue une recherche avancée floue (nom, propriétaire, ville).")
    add_actor(doc, "Systèmes externes",
        "Google OAuth (authentification), OpenStreetMap/Nominatim (géocodage), stockage fichiers local.")

    add_placeholder(doc, "FIGURE 2.1 — Diagramme des cas d'utilisation",
                    "Diagramme UML use case avec les 4 acteurs principaux et leurs cas d'utilisation.", 2)

    add_heading(doc, "2.2 Besoins fonctionnels", 2)

    add_heading(doc, "2.2.1 Module Authentification", 3)
    for f in [
        "Inscription et connexion client/propriétaire (email/mot de passe).",
        "Connexion Google OAuth 2.0 (Google Sign-In).",
        "Connexion administrateur par identifiants dédiés.",
        "Gestion du profil : modification, changement mot de passe, suppression de compte.",
        "Déconnexion sécurisée avec invalidation du token local.",
    ]:
        add_bullet(doc, f)

    add_heading(doc, "2.2.2 Module Hôtels", 3)
    for f in [
        "Ajout d'hôtel par propriétaire (nom, ville, adresse, description, GPS).",
        "Soumission en statut en_attente pour validation admin.",
        "Validation/refus admin avec notification automatique.",
        "Modification/suppression avec vérification des réservations actives.",
        "Upload image principale et galerie photos.",
        "Recherche publique (ville, type, capacité, budget) et hôtels populaires.",
        "Recherche admin avancée floue et recherche propriétaire fusionnée au bandeau d'accueil.",
    ]:
        add_bullet(doc, f)

    add_heading(doc, "2.2.3 Module Chambres", 3)
    for f in [
        "CRUD chambres (type, prix, capacité) par propriétaire.",
        "Types : Économique, Standard, Supérieure, Deluxe, Suite, Familiale.",
        "Galerie photos par chambre ; suppression conditionnée aux réservations actives.",
    ]:
        add_bullet(doc, f)

    add_heading(doc, "2.2.4 Module Réservations", 3)
    for f in [
        "Réservation client avec sélection des dates.",
        "Vérification automatique des chevauchements de dates.",
        "Workflow : en_attente → confirmée / refusée (propriétaire) ou annulée (client).",
        "Modification client uniquement si statut en_attente.",
        "Historique client et vue propriétaire consolidée.",
    ]:
        add_bullet(doc, f)

    add_heading(doc, "2.2.5 Module Avis & Notes", 3)
    for f in [
        "Avis (note 1–5 + commentaire) par client et par hôtel (unique, modifiable).",
        "Affichage moyenne et nombre d'avis sur chaque fiche.",
    ]:
        add_bullet(doc, f)

    add_heading(doc, "2.2.6 Module Favoris", 3)
    add_bullet(doc, "Ajout/retrait d'hôtels en favoris ; liste avec prix minimum et date d'ajout.")

    add_heading(doc, "2.2.7 Module Notifications", 3)
    for f in [
        "Notifications in-app : inscription, soumission/validation/refus hôtel, réservation, "
        "confirmation/refus/annulation, nouvel avis.",
        "Compteur non lues (notificationsBell.js) ; marquage individuel ou global comme lu.",
    ]:
        add_bullet(doc, f)

    add_heading(doc, "2.2.8 Module Statistiques", 3)
    for f in [
        "Dashboard propriétaire : hôtels, chambres, réservations, revenus, répartition type/capacité.",
        "Dashboard admin : totaux globaux (clients, propriétaires, hôtels par statut).",
    ]:
        add_bullet(doc, f)

    add_heading(doc, "2.3 Besoins non fonctionnels", 2)
    nfunc = [
        ("Performance", "Réponses API < 500 ms en conditions normales ; cache sessionStorage accueil ; limite JSON 12 Mo."),
        ("Sécurité", "JWT (2 h), bcrypt (10 rounds), contrôle rôle, validation uploads (MIME, 5 Mo max)."),
        ("Disponibilité", "Architecture monolithique simple à déployer ; erreurs centralisées côté Express."),
        ("Ergonomie", "Charte #2F64EA / #FBBF24, Plus Jakarta Sans, navigation cohérente, responsive."),
        ("Maintenabilité", "Routes modulaires, migrations SQL incrémentales, frontend séparé par rôle."),
        ("Scalabilité", "REST + séparation frontend/backend ; évolution vers pool MySQL et pagination."),
        ("Compatibilité", "Chrome, Firefox, Edge, Safari ; design responsive mobile/tablette."),
    ]
    for name, desc in nfunc:
        p = doc.add_paragraph()
        p.add_run(f"{name} : ").bold = True
        p.add_run(desc)

    add_heading(doc, "2.4 User Stories", 2)
    add_user_story(doc, "US-01", "visiteur souhaitant créer un compte client",
        "m'inscrire avec nom, prénom, email et mot de passe",
        "accéder aux fonctionnalités de réservation et de gestion de profil")
    add_user_story(doc, "US-02", "utilisateur ayant un compte Google",
        "me connecter en un clic via Google OAuth",
        "éviter de créer un nouveau mot de passe et accéder rapidement")
    add_user_story(doc, "US-03", "client planifiant un voyage à Marrakech",
        "rechercher des hôtels filtrés par type et budget maximum",
        "trouver un hébergement adapté à mes besoins")
    add_user_story(doc, "US-04", "client ayant trouvé un hôtel",
        "réserver une chambre en choisissant mes dates d'arrivée et de départ",
        "sécuriser mon hébergement pour mon séjour")
    add_user_story(doc, "US-05", "propriétaire d'hôtel",
        "consulter et accepter ou refuser les réservations reçues",
        "contrôler le planning d'occupation de mon établissement")
    add_user_story(doc, "US-06", "administrateur de la plateforme",
        "valider ou refuser les hôtels soumis par les propriétaires",
        "garantir la qualité des offres présentées aux clients")
    add_user_story(doc, "US-07", "client ayant séjourné dans un hôtel",
        "laisser une note et un commentaire",
        "partager mon expérience et aider d'autres voyageurs")
    add_user_story(doc, "US-08", "propriétaire d'hôtel",
        "accéder à un tableau de bord (revenus, réservations, performances chambres)",
        "suivre l'activité et prendre des décisions éclairées")

    add_heading(doc, "2.5 Analyse des données", 2)
    add_body(doc,
        "Les données sont principalement relationnelles (utilisateurs, hôtels, chambres, "
        "réservations, images, favoris, notifications, avis) complétées par des fichiers "
        "binaires (uploads/hotels/, uploads/chambres/).")
    add_bullet(doc, "Type : structuré (SQL) + fichiers binaires.")
    add_bullet(doc, "Source : saisie utilisateur, uploads, événements métier.")
    add_bullet(doc, "Volume : prototype académique ; scalable via index et pagination future.")

    add_placeholder(doc, "TABLEAU 2.1 — Matrice acteurs × fonctionnalités",
                    "Tableau croisé : Visiteur, Client, Propriétaire, Admin × modules.", 1)
    doc.add_page_break()

    # ===================== CHAPITRE 3 =====================
    add_heading(doc, "Chapitre 3 : Conception", 1)

    add_heading(doc, "3.1 Architecture globale", 2)
    add_body(doc, "Architecture three-tier (client-serveur) :")
    add_bullet(doc, "Présentation : frontend HTML/CSS/JS (Public, Client, Propriétaire, Admin, Commun).")
    add_bullet(doc, "Métier : API REST Express (backend/server.js, 10 routeurs modulaires).")
    add_bullet(doc, "Données : MySQL (gestion_hotel) + système de fichiers pour les médias.")
    add_body(doc,
        "Express sert l'API (/api) et les fichiers statiques du frontend (port 3000). "
        "Les échanges se font en JSON via HTTP.")

    add_placeholder(doc, "FIGURE 3.1 — Architecture 3-tiers",
                    "Schéma : Navigateur ↔ Express.js API REST ↔ MySQL / Uploads.", 2)

    add_heading(doc, "3.2 Choix technologiques", 2)
    add_placeholder(doc, "LOGOS TECHNO",
                    "Logos : Node.js, Express, MySQL, JavaScript, HTML5, CSS3, JWT, Leaflet, Google OAuth.", 2)

    add_heading(doc, "3.2.1 Backend", 3)
    add_body(doc,
        "Node.js pour les performances I/O et l'écosystème npm. Express.js 5 pour structurer "
        "l'API REST de manière modulaire et légère.")
    add_heading(doc, "3.2.2 Base de données", 3)
    add_body(doc,
        "MySQL pour les garanties ACID, l'intégrité référentielle et les contraintes CHECK. "
        "Driver mysql2 avec support Promises.")
    add_heading(doc, "3.2.3 Frontend", 3)
    add_body(doc,
        "HTML/CSS/JavaScript vanilla : maîtrise totale de la pile, pas de build complexe, "
        "portails distincts par rôle. Feuille principale style.css (~13 000 lignes).")
    add_heading(doc, "3.2.4 Authentification et sécurité", 3)
    add_body(doc,
        "JWT stateless (2 h), Google OAuth 2.0 vérifié côté serveur (tokeninfo), bcrypt pour les mots de passe.")

    add_table(doc,
        ["Catégorie", "Technologie", "Version", "Rôle"],
        [
            ("Runtime", "Node.js", "LTS", "Exécution JavaScript serveur"),
            ("Framework", "Express.js", "5.x", "API REST, middlewares"),
            ("Base de données", "MySQL", "8.x", "Stockage persistant"),
            ("Driver DB", "mysql2", "3.x", "Connexion Node ↔ MySQL"),
            ("Auth", "jsonwebtoken", "9.x", "JWT"),
            ("Sécurité", "bcrypt", "6.x", "Hachage mots de passe"),
            ("Auth sociale", "Google OAuth 2.0", "—", "Connexion Google"),
            ("Cartographie", "Leaflet + OSM", "—", "Géolocalisation hôtels"),
            ("Frontend", "HTML/CSS/JS", "Vanilla", "Interfaces multi-portail"),
        ])

    add_heading(doc, "3.3 Diagrammes de conception", 2)
    add_body(doc, "Diagrammes à insérer :")
    add_bullet(doc, "Classes : Utilisateur, Hôtel, Chambre, Réservation, Avis, Notification, Favori.")
    add_bullet(doc, "Séquence — Réservation client → validation propriétaire.")
    add_bullet(doc, "Séquence — Soumission hôtel → modération admin.")
    add_bullet(doc, "Déploiement — Serveur Node, MySQL, dossier uploads.")

    add_placeholder(doc, "FIGURE 3.2 — Diagramme de classes UML",
                    "Diagramme UML complet des entités métier.", 3)
    add_placeholder(doc, "FIGURE 3.3 — Diagramme de séquence (réservation)",
                    "Client → Recherche → Chambre → Réservation → Notification propriétaire.", 3)
    add_placeholder(doc, "FIGURE 3.4 — Diagramme de séquence (validation hôtel)",
                    "Propriétaire → Soumission → Admin validation → Notification.", 2)

    add_heading(doc, "3.4 Base de données", 2)
    add_body(doc, "La base gestion_hotel comprend 9 tables principales :")
    add_table(doc,
        ["Table", "Colonnes clés", "Description"],
        [
            ("utilisateurs", "id, nom, prenom, email, mot_de_passe, role, google_id", "Clients, propriétaires, admins"),
            ("hotels", "id, nom, ville, adresse, statut, latitude, longitude, id_proprietaire", "Établissements hôteliers"),
            ("chambres", "id, id_hotel, type, prix, capacite", "Chambres rattachées à un hôtel"),
            ("reservations", "id, id_chambre, id_client, date_debut, date_fin, statut", "Réservations avec workflow"),
            ("hotel_images", "id, id_hotel, chemin, ordre", "Galerie photos hôtels"),
            ("chambre_images", "id, id_chambre, chemin, ordre", "Galerie photos chambres"),
            ("avis", "id, id_client, id_hotel, note, commentaire", "Notes et commentaires (UNIQUE client/hôtel)"),
            ("favoris", "id, id_client, id_hotel, created_at", "Hôtels favoris des clients"),
            ("notifications", "id, id_utilisateur, type, titre, message, lien, lue", "Notifications in-app"),
        ])
    add_body(doc, "Relations : propriétaire 1→N hôtels ; hôtel 1→N chambres ; chambre 1→N réservations "
             "(avec contrôle chevauchement) ; client 1→N réservations ; avis unique par client/hôtel.")

    add_placeholder(doc, "FIGURE 3.5 — Schéma relationnel (MLD / MPD)",
                    "Diagramme entité-association ou MPD MySQL complet.", 3)

    add_heading(doc, "3.5 Conception des APIs REST", 2)
    add_body(doc, "Base URL : http://localhost:3000/api")

    add_heading(doc, "3.5.1 Module Authentification", 3)
    add_table(doc, ["Endpoint", "Méthode", "Rôle", "Description"], [
        ("/auth/client/register", "POST", "Public", "Inscription client"),
        ("/auth/client/login", "POST", "Public", "Connexion client"),
        ("/auth/proprietaire/register", "POST", "Public", "Inscription propriétaire"),
        ("/auth/proprietaire/login", "POST", "Public", "Connexion propriétaire"),
        ("/auth/admin/login", "POST", "Public", "Connexion admin"),
        ("/auth/google", "POST", "Public", "OAuth Google"),
        ("/profil", "GET/PUT/DELETE", "Authentifié", "Profil utilisateur"),
        ("/profil/password", "PUT", "Authentifié", "Changement mot de passe"),
    ])

    add_heading(doc, "3.5.2 Module Hôtels", 3)
    add_table(doc, ["Endpoint", "Méthode", "Rôle", "Description"], [
        ("/hotels/validated", "GET", "Public", "Hôtels validés"),
        ("/hotels/search", "GET", "Public", "Recherche multi-critères"),
        ("/hotels/popular", "GET", "Public", "Top hôtels populaires"),
        ("/villes/popular", "GET", "Public", "Top villes"),
        ("/hotels", "POST", "Propriétaire", "Créer un hôtel (en_attente)"),
        ("/hotels/mes-hotels", "GET", "Propriétaire", "Mes hôtels"),
        ("/hotels/:id", "PUT/DELETE", "Propriétaire", "Modifier/supprimer"),
        ("/hotels/admin/list", "GET", "Admin", "Liste filtrée admin"),
        ("/hotels/:id/valider", "PUT", "Admin", "Valider un hôtel"),
        ("/hotels/:id/refuser", "PUT", "Admin", "Refuser un hôtel"),
    ])

    add_heading(doc, "3.5.3 Module Réservations", 3)
    add_table(doc, ["Endpoint", "Méthode", "Rôle", "Description"], [
        ("/reservations", "POST", "Client", "Créer réservation"),
        ("/reservations/:id", "PUT", "Client", "Modifier (si en_attente)"),
        ("/reservations/:id/annuler", "PUT", "Client", "Annuler"),
        ("/reservations/mes-reservations", "GET", "Client", "Historique client"),
        ("/reservations/proprietaire/detail", "GET", "Propriétaire", "Réservations reçues"),
        ("/reservations/:id/valider", "PUT", "Propriétaire", "Confirmer"),
        ("/reservations/:id/refuser", "PUT", "Propriétaire", "Refuser"),
    ])

    add_heading(doc, "3.5.4 Autres modules", 3)
    add_table(doc, ["Endpoint", "Méthode", "Rôle", "Description"], [
        ("/avis", "POST/DELETE", "Client", "Publier/supprimer avis"),
        ("/hotels/:id/avis", "GET", "Public", "Lire avis + moyenne"),
        ("/favoris", "GET", "Client", "Lister favoris"),
        ("/favoris/:hotelId", "POST/DELETE", "Client", "Ajouter/retirer favori"),
        ("/notifications", "GET", "Authentifié", "Lister notifications"),
        ("/notifications/lire-tout", "PUT", "Authentifié", "Tout marquer lu"),
        ("/proprietaire/stats", "GET", "Propriétaire", "Stats propriétaire"),
        ("/admin/stats", "GET", "Admin", "Stats globales"),
    ])

    add_heading(doc, "3.6 Conception du module intelligent", 2)
    add_body(doc,
        "Le module « Smart » repose sur des heuristiques métier (pas de ML pour l'instant) :")
    for item in [
        "Classement hôtels/villes populaires (réservations confirmées).",
        "Recherche floue admin : normalisation NFD, suppression accents, filtres combinés.",
        "Recherche propriétaire fusionnée au hero de hotel.html.",
        "Politique de suppression intelligente (deletionPolicy.js).",
        "Notifications contextuelles (notificationService.js, 10+ types).",
        "Vérification chevauchement dates de réservation.",
    ]:
        add_bullet(doc, item)
    add_body(doc,
        "Évolution IA envisagée : recommandation collaborative, prédiction d'occupation, "
        "détection d'anomalies sur les avis.")

    add_placeholder(doc, "FIGURE 3.6 — Flux notifications intelligentes",
                    "Schéma notificationService → base → notificationsBell.js frontend.", 2)
    doc.add_page_break()

    # ===================== CHAPITRE 4 =====================
    add_heading(doc, "Chapitre 4 : Développement", 1)

    add_heading(doc, "4.1 Environnement de développement", 2)
    add_table(doc, ["Outil", "Utilisation"], [
        ("Visual Studio Code / Cursor", "Éditeur principal"),
        ("Node.js (LTS)", "Runtime backend"),
        ("MySQL Workbench", "Administration base de données"),
        ("Postman / REST Client", "Tests API"),
        ("Git & GitHub", "Versionnement et collaboration"),
        ("Chrome DevTools", "Débogage frontend"),
        ("npm", "Gestion dépendances"),
    ])

    add_heading(doc, "4.2 Structure et implémentation", 2)
    add_body(doc, "Organisation du projet gestion-hotel/ :")
    structure = [
        ("backend/server.js", "Point d'entrée Express, montage routes et statiques."),
        ("backend/routes/", "10 routeurs : auth, hotels, chambres, réservations, stats, favoris, notifications, avis, images."),
        ("backend/middleware/authMiddleware.js", "Vérification JWT Bearer."),
        ("backend/utils/", "notificationService, deletionPolicy, hotelImageStorage."),
        ("backend/database/", "Migrations SQL incrémentales."),
        ("frontend/Public/", "Vitrine : index, recherche, détail hôtel."),
        ("frontend/Client/", "Dashboard, réservation, favoris, avis."),
        ("frontend/Proprietaire/", "Hôtels, chambres, réservations, stats, recherche fusionnée."),
        ("frontend/Admin/", "Modération, recherche avancée floue, statistiques."),
        ("frontend/Commun/", "api.js, auth, notificationsBell, hotelMap, homeData, appBrand."),
        ("frontend/css/style.css", "Charte #2F64EA / #FBBF24, responsive."),
    ]
    for path, desc in structure:
        p = doc.add_paragraph()
        p.add_run(f"{path} — ").bold = True
        p.add_run(desc)

    add_heading(doc, "4.2.1 Module Authentification", 3)
    add_body(doc,
        "Connexion classique : bcrypt compare le mot de passe, JWT signé (id, rôle, 2 h). "
        "Google OAuth : credential vérifié via API tokeninfo, création/mise à jour utilisateur. "
        "Routes séparées par portail avec messages explicites en cas de mauvais rôle.")
    add_placeholder(doc, "CAPTURE 4.1 — Connexion Client",
                    "Capture login.html client + bouton Google OAuth.", 1)

    add_heading(doc, "4.2.2 Module Hôtels et Chambres", 3)
    add_body(doc,
        "Workflow hôtel : en_attente → valide / refuse. Visibilité publique uniquement si valide. "
        "GPS validé côté serveur (±90°/±180°). deletionPolicy.js bloque suppression si réservations actives.")
    add_placeholder(doc, "CAPTURE 4.2 — Dashboard Propriétaire",
                    "Liste hôtels avec statuts + formulaire hotelAjouter.html avec carte Leaflet.", 2)

    add_heading(doc, "4.2.3 Module Réservation", 3)
    add_body(doc,
        "Contrôle disponibilité par intersection d'intervalles SQL. Statuts : en_attente → "
        "confirmee / refusee / annulee. Modification client seulement si en_attente.")
    add_placeholder(doc, "CAPTURE 4.3 — Réservation",
                    "Interface réservation client + gestion réservations propriétaire.", 2)

    add_heading(doc, "4.2.4 Module Notifications", 3)
    add_body(doc,
        "notificationService.js : fonction fire() asynchrone non bloquante. Types : bienvenue_client, "
        "hotel_en_attente, hotel_valide, hotel_refuse, reservation_en_attente, reservation_confirmee, "
        "reservation_refusee, reservation_annulee, avis_nouveau. notificationsBell.js interroge l'API périodiquement.")
    add_placeholder(doc, "CAPTURE 4.4 — Notifications",
                    "Cloche avec compteur et liste déroulante notifications non lues.", 1)

    add_heading(doc, "4.2.5 Module Avis et Favoris", 3)
    add_body(doc,
        "Avis : contrainte UNIQUE (id_client, id_hotel) + ON DUPLICATE KEY UPDATE. "
        "Favoris : liste personnalisée avec prix minimum et date d'ajout.")
    add_placeholder(doc, "CAPTURE 4.5 — Avis et Favoris",
                    "Fiche hôtel avec étoiles avisStars.js + page favoris client.", 1)

    add_heading(doc, "4.2.6 Tableaux de bord et recherche admin", 3)
    add_body(doc,
        "Dashboard propriétaire : requête SQL agrégée (hôtels, chambres, réservations, revenus). "
        "Dashboard admin : totaux globaux. Recherche admin (rechercheResultat.js) : normalisation "
        "accents, filtres combinés. Recherche propriétaire fusionnée au bandeau hero de hotel.html.")
    add_placeholder(doc, "CAPTURE 4.6 — Dashboards et recherche",
                    "Stats propriétaire/admin + résultats recherche admin floue.", 2)

    add_heading(doc, "4.3 Interfaces utilisateurs", 2)
    add_placeholder(doc, "CAPTURE 4.7 — Accueil public",
                    "Hero, barre recherche, hôtels populaires, villes populaires (index.html).", 1)
    add_placeholder(doc, "CAPTURE 4.8 — Résultats recherche",
                    "Page rechercheResultat avec filtres et cartes hôtels.", 1)
    add_placeholder(doc, "CAPTURE 4.9 — Fiche hôtel détaillée",
                    "Galerie, description, carte Leaflet, chambres, avis.", 1)

    add_heading(doc, "4.4 Gestion de la sécurité", 2)
    add_heading(doc, "4.4.1 Authentification et autorisation", 3)
    add_body(doc,
        "Header Authorization: Bearer <token>. authMiddleware.js injecte req.user. "
        "Ségrégation stricte : propriétaire → ses ressources ; client → ses réservations/avis ; admin → routes admin.")
    add_heading(doc, "4.4.2 Protection des données", 3)
    for s in [
        "Mots de passe hashés bcrypt (10 rounds), jamais renvoyés en API.",
        "Emails normalisés (trim + toLowerCase).",
        "Suppression compte : vérification réservations/hôtels actifs.",
        "Tokens Google vérifiés côté serveur uniquement.",
        "Validation uploads : MIME autorisés, 5 Mo max.",
    ]:
        add_bullet(doc, s)
    add_heading(doc, "4.4.3 Validation des données", 3)
    for s in [
        "Champs obligatoires vérifiés (nom, ville, dates).",
        "GPS dans plages légales ; 8 villes marocaines autorisées.",
        "Notes avis 1–5 (CHECK SQL) ; commentaires 5–2000 caractères.",
    ]:
        add_bullet(doc, s)

    add_heading(doc, "4.5 Difficultés rencontrées", 2)
    add_heading(doc, "4.5.1 Techniques", 3)
    for d in [
        "Chevauchements dates réservation : requête SQL itérative.",
        "deletionPolicy.js : contraintes métier avant suppression.",
        "Upload images : cohérence base/disque (hotels/{id}/, chambres/{id}/).",
        "Google OAuth sans lib dédiée : gestion tokeninfo et cas limites.",
        "CSS unifié multi-portail sans framework (~13 000 lignes).",
        "Recherche admin floue + fusion recherche propriétaire.",
    ]:
        add_bullet(doc, d)
    add_heading(doc, "4.5.2 Organisationnelles", 3)
    for d in [
        "Priorisation MVP dans un périmètre large.",
        "Coordination Git et résolution conflits.",
        "Tests manuels Postman + interfaces (pas de suite automatisée).",
    ]:
        add_bullet(doc, d)

    add_heading(doc, "4.6 Tests effectués", 2)
    for t in [
        "Inscription/connexion (client, propriétaire, admin, Google).",
        "Création hôtel → modération → visibilité publique.",
        "Recherche multi-critères et réservation bout en bout.",
        "Validation/refus réservation ; favoris ; avis ; notifications.",
        "Upload/suppression images ; responsive mobile/tablette.",
        "Recherche admin floue et recherche propriétaire fusionnée.",
    ]:
        add_bullet(doc, t)

    add_placeholder(doc, "TABLEAU 4.1 — Plan de tests fonctionnels",
                    "Scénario | étapes | résultat attendu | statut.", 2)
    doc.add_page_break()

    # ===================== CHAPITRE 5 =====================
    add_heading(doc, "Chapitre 5 : Résultats et perspectives", 1)

    add_heading(doc, "5.1 Stratégie de test", 2)
    add_heading(doc, "5.1.1 Tests API", 3)
    add_body(doc, "Tests manuels Postman/REST Client pour chaque endpoint :")
    for t in ["Cas nominal (200/201)", "Validation (400)", "Non autorisé (401/403)",
              "Conflit métier (409)", "Introuvable (404)"]:
        add_bullet(doc, t)
    add_heading(doc, "5.1.2 Tests fonctionnels E2E", 3)
    for t in [
        "Réservation complète : inscription → recherche → réservation → validation → notification.",
        "Publication hôtel : propriétaire → admin → visibilité publique.",
        "Suppression bloquée si réservations actives.",
        "Connexion Google : création compte puis reconnexions.",
    ]:
        add_bullet(doc, t)

    add_heading(doc, "5.2 Résultats", 2)
    add_heading(doc, "5.2.1 Fonctionnalités implémentées", 3)
    add_table(doc, ["Fonctionnalité", "Statut"], [
        ("Authentification JWT + Google OAuth", "Réalisée"),
        ("Gestion hôtels/chambres + modération admin", "Réalisée"),
        ("Galerie photos + géolocalisation Leaflet", "Réalisée"),
        ("Recherche filtrée + hôtels populaires", "Réalisée"),
        ("Recherche admin floue + fusion propriétaire", "Réalisée"),
        ("Réservation avec contrôle disponibilités", "Réalisée"),
        ("Avis, favoris, notifications temps réel", "Réalisée"),
        ("Dashboards statistiques admin/propriétaire", "Réalisée"),
        ("Politique suppression intelligente", "Réalisée"),
        ("Interface publique responsive", "Réalisée"),
        ("Module IA prédictif (ML)", "Non implémenté — perspective"),
    ])

    add_heading(doc, "5.2.2 Analyse", 3)
    add_body(doc,
        "L'ensemble du cahier des charges est couvert. Performance locale satisfaisante "
        "(< 100 ms opérations simples, < 300 ms requêtes agrégées). Navigation admin et "
        "propriétaire clarifiée après refonte headers ; recherche fusionnée améliore l'ergonomie.")

    add_heading(doc, "5.2.3 Limites", 3)
    for l in [
        "Pas de tests automatisés (Jest/Mocha).",
        "Connexion MySQL simple (non poolée).",
        "Pas de pagination sur certains listings.",
        "Stockage images local (pas de CDN).",
        "Pas de paiement en ligne.",
        "8 villes marocaines seulement.",
        "Module ML non implémenté (heuristiques uniquement).",
    ]:
        add_bullet(doc, l)

    add_heading(doc, "5.2.4 Comparaison avec l'existant", 3)
    add_body(doc,
        "Contrairement aux OTA internationales, HôtelFacile Smart offre contrôle total du code, "
        "modération intégrée, gratuité et adaptation au contexte marocain. Compromis : catalogue "
        "limité au périmètre académique.")

    add_heading(doc, "5.3 Améliorations futures", 2)
    add_heading(doc, "5.3.1 Court terme", 3)
    for f in ["Pool MySQL", "Pagination serveur", "Tests Jest", "Validation frontend renforcée"]:
        add_bullet(doc, f)
    add_heading(doc, "5.3.2 Moyen terme", 3)
    for f in ["Paiement CMI/Stripe", "App mobile/PWA", "Messagerie interne", "Stockage cloud (S3)", "Module IA recommandation"]:
        add_bullet(doc, f)
    add_heading(doc, "5.3.3 Long terme", 3)
    for f in ["Extension Maghreb", "Intégration PMS", "Analytics avancés", "Programme fidélité"]:
        add_bullet(doc, f)

    add_placeholder(doc, "FIGURE 5.1 — Graphiques de résultats",
                    "Répartition statuts hôtels, réservations par mois, notes moyennes.", 2)
    doc.add_page_break()

    # ===================== CONCLUSION =====================
    add_heading(doc, "Conclusion générale", 1)

    add_heading(doc, "Résumé du projet", 2)
    add_body(doc,
        "Ce projet nous a permis de concevoir et développer HôtelFacile Smart de A à Z : "
        "une plateforme web de gestion hôtelière couvrant l'inscription propriétaire, la "
        "validation admin, la réservation client, la gestion quotidienne et les avis post-séjour.")
    add_body(doc,
        "L'architecture Node.js/Express + MySQL + frontend vanilla s'est révélée adaptée : "
        "simplicité de déploiement, contrôle total et performances suffisantes.")

    add_heading(doc, "Apports du projet", 2)
    add_bullet(doc, "Modélisation UML et base relationnelle complexe.")
    add_bullet(doc, "API REST sécurisée (JWT, OAuth, RBAC).")
    add_bullet(doc, "Workflows multi-états (hôtels, réservations).")
    add_bullet(doc, "UX/UI cohérente (#2F64EA / #FBBF24) sur 4 portails.")

    add_heading(doc, "Compétences acquises", 2)
    for s in [
        "Backend Node.js/Express, middlewares, gestion erreurs.",
        "SQL : migrations, jointures, sous-requêtes, contraintes.",
        "Sécurité web : JWT, OAuth, bcrypt, RBAC.",
        "Frontend vanilla : DOM, fetch, localStorage, responsive CSS.",
        "Gestion de projet itérative et travail collaboratif Git.",
    ]:
        add_bullet(doc, s)

    add_heading(doc, "Mot de fin", 2)
    add_body(doc,
        "Ce projet dépasse l'exercice académique : il répond à un besoin réel du tourisme "
        "marocain. HôtelFacile Smart est fonctionnelle, sécurisée et prête pour un déploiement "
        "pilote. Nous espérons qu'elle servira de base solide à la digitalisation des "
        "hôteliers indépendants au Maroc.")
    doc.add_page_break()

    # ===================== BIBLIOGRAPHIE =====================
    add_heading(doc, "Bibliographie", 1)
    add_heading(doc, "Documentation technique", 2)
    refs_tech = [
        "Node.js Documentation — https://nodejs.org/en/docs/",
        "Express.js Documentation — https://expressjs.com/",
        "MySQL 8.0 Reference Manual — https://dev.mysql.com/doc/refman/8.0/en/",
        "JSON Web Token (JWT) — RFC 7519 — https://datatracker.ietf.org/doc/html/rfc7519",
        "Google Identity Services — https://developers.google.com/identity",
        "Leaflet — https://leafletjs.com/",
        "OpenStreetMap / Nominatim — https://nominatim.org/",
        "OWASP REST Security Cheat Sheet — https://cheatsheetseries.owasp.org/",
        "MDN Web Docs — https://developer.mozilla.org/",
    ]
    for i, ref in enumerate(refs_tech, 1):
        add_body(doc, f"[{i}] {ref}")

    add_heading(doc, "Articles et ouvrages", 2)
    refs_other = [
        "Ministère du Tourisme du Maroc — https://www.tourisme.gov.ma",
        "Jon Duckett — HTML & CSS: Design and Build Websites — Wiley, 2011",
        "Thomas Connolly & Carolyn Begg — Database Systems — Pearson, 2014",
        "Booking.com, Airbnb, Expedia — Analyse comparative OTA (2024–2025)",
    ]
    for i, ref in enumerate(refs_other, len(refs_tech) + 1):
        add_body(doc, f"[{i}] {ref}")

    doc.add_page_break()

    # ===================== ANNEXES =====================
    add_heading(doc, "Annexes", 1)

    add_heading(doc, "Annexe A — Arborescence du projet", 2)
    add_body(doc,
        "gestion-hotel/\n"
        "├── backend/ (config, database, middleware, routes, utils, uploads, server.js)\n"
        "└── frontend/ (Public, Client, Proprietaire, Admin, Commun, css, images)")

    add_heading(doc, "Annexe B — Diagrammes UML", 2)
    add_placeholder(doc, "DIAGRAMMES UML COMPLETS",
                    "Use case, classes, séquences, déploiement, MLD, diagrammes d'activité.", 4)

    add_heading(doc, "Annexe C — Captures d'écran", 2)
    add_placeholder(doc, "CAPTURES COMPLÉMENTAIRES",
                    "Login, register, détail chambre, profil, notifications, recherche admin/propriétaire.", 4)

    add_heading(doc, "Annexe D — Extrait server.js", 2)
    add_body(doc,
        "require('./config/env');\n"
        "const express = require('express');\n"
        "app.use(express.json({ limit: '12mb' }));\n"
        "app.use('/api', require('./routes/authRoutes'));\n"
        "app.use('/api', require('./routes/hotelRoutes'));\n"
        "// ... 8 autres routeurs\n"
        "app.listen(PORT, () => console.log(`Serveur lancé sur http://localhost:${PORT}`));")

    add_heading(doc, "Annexe E — Migration SQL (avis)", 2)
    add_body(doc,
        "CREATE TABLE avis (\n"
        "  id INT AUTO_INCREMENT PRIMARY KEY,\n"
        "  note TINYINT NOT NULL CHECK (note >= 1 AND note <= 5),\n"
        "  id_client INT NOT NULL, id_hotel INT NOT NULL,\n"
        "  UNIQUE KEY uk_avis_client_hotel (id_client, id_hotel)\n"
        ");")

    add_heading(doc, "Annexe F — Variables d'environnement", 2)
    add_body(doc, "PORT=3000 | JWT_SECRET=... | GOOGLE_CLIENT_ID=... | MySQL: gestion_hotel")

    add_placeholder(doc, "LOGOS TECHNO (ANNEXE)",
                    "Planche : Node.js, Express, MySQL, JWT, Leaflet, Google, HTML5, CSS3, JavaScript.", 2)

    doc.save(OUTPUT)
    print(f"Rapport généré : {OUTPUT}")


if __name__ == "__main__":
    build()
