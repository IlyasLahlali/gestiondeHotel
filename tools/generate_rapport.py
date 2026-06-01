# -*- coding: utf-8 -*-
"""Génère le rapport de projet HôtelFacile Smart dans docs/Rapport.docx"""

import os

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r"C:\Users\ilyas\Desktop\gestion-hotel\docs\Rapport.docx"
OUTPUT_TMP = r"C:\Users\ilyas\Desktop\gestion-hotel\docs\Rapport_generated.docx"

BLUE = RGBColor(0x2F, 0x64, 0xEA)
GOLD = RGBColor(0xFB, 0xBF, 0x24)
DARK = RGBColor(0x11, 0x18, 0x27)
GRAY = RGBColor(0x64, 0x74, 0x8B)


def set_cell_shading(cell, fill_hex: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_placeholder(doc, title, hint, height_lines=3):
    p = doc.add_paragraph()
    run = p.add_run(f"【 {title} 】")
    run.bold = True
    run.font.color.rgb = BLUE
    for _ in range(height_lines):
        box = doc.add_paragraph()
        box.add_run(hint).italic = True
        box.paragraph_format.left_indent = Cm(1)
        box.paragraph_format.space_after = Pt(6)
    doc.add_paragraph()


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = BLUE if level == 1 else DARK
    return h


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Cm(1 + level * 0.5)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_table_api(doc, rows):
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    headers = ["Endpoint", "Méthode", "Rôle", "Description"]
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], "EFF6FF")
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph()


def add_table_generic(doc, headers, rows, fill="EFF6FF"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], fill)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph()


def collect_project_files(root):
    skip = {".git", "node_modules", "__pycache__", ".cursor"}
    entries = []
    for dirpath, dirnames, filenames in os.walk(root):
        dirnames[:] = [d for d in dirnames if d not in skip]
        rel = os.path.relpath(dirpath, root)
        for fn in sorted(filenames):
            if fn.startswith("."):
                continue
            path = fn if rel == "." else os.path.join(rel, fn)
            entries.append(path.replace("\\", "/"))
    return sorted(entries)


def build():
    doc = Document()
    sections = doc.sections
    for s in sections:
        s.top_margin = Cm(2.5)
        s.bottom_margin = Cm(2.5)
        s.left_margin = Cm(2.5)
        s.right_margin = Cm(2.5)

    # Styles de base
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # ===================== PAGE DE GARDE =====================
    for _ in range(6):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("HôtelFacile Smart")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = BLUE

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("Plateforme web intelligente de gestion et de réservation d'hôtels au Maroc")
    rs.font.size = Pt(14)
    rs.font.color.rgb = DARK

    doc.add_paragraph()
    p_upm = doc.add_paragraph()
    p_upm.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_upm = p_upm.add_run("Université Privée de Marrakech (UPM)")
    r_upm.italic = True
    r_upm.font.size = Pt(12)
    r_upm.font.color.rgb = GRAY

    info = [
        ("Filière", "Développement Logiciel & Sécurité Informatique"),
        ("Spécialité", "Développement Logiciel & Sécurité Informatique"),
        ("Thématique", "Tourisme / Hôtellerie numérique"),
        ("Réalisé par", "[À compléter]\n[À compléter]"),
        ("Encadré par", "[À compléter]"),
        ("Établissement", "[À compléter]"),
        ("Lieu et date", "Marrakech — 2026"),
        ("Année universitaire", "2025 – 2026"),
    ]
    for label, value in info:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"{label} : ").bold = True
        p.add_run(value)

    add_placeholder(doc, "LOGO ÉTABLISSEMENT / LOGO PROJET",
                    "Insérer ici le logo de l'établissement et/ou le logo HôtelFacile Smart (frontend/images/logo.svg).", 1)
    add_placeholder(doc, "LOGOS TECHNO",
                    "Planche des technologies : Node.js, Express, MySQL, HTML5, CSS3, JavaScript, JWT, Leaflet, Google OAuth.", 1)
    doc.add_page_break()

    # ===================== REMERCIEMENTS =====================
    add_heading(doc, "Remerciements", 1)
    add_body(doc,
        "Nous tenons en premier lieu à exprimer notre profonde gratitude envers notre encadrant, dont la "
        "disponibilité, la rigueur et les précieux conseils ont guidé notre travail tout au long de ce "
        "projet. Son accompagnement bienveillant a été une source constante de motivation et d'inspiration.")
    add_body(doc,
        "Nous remercions également l'ensemble du corps enseignant de l'UPM et du parcours Développement "
        "Logiciel & Sécurité Informatique pour la formation de qualité dispensée au cours de ces années "
        "d'études, ainsi que pour les valeurs académiques et professionnelles transmises.")
    add_body(doc,
        "Nos remerciements vont aussi à nos familles et à nos proches pour leur soutien indéfectible, "
        "leur patience et leurs encouragements tout au long de ce parcours.")
    add_body(doc,
        "Enfin, nous remercions tous ceux qui, de près ou de loin, ont contribué à la réalisation de "
        "HôtelFacile Smart, que ce soit par leurs retours, leurs idées ou leur aide technique.")
    doc.add_page_break()

    # ===================== RÉSUMÉ =====================
    add_heading(doc, "Résumé (Abstract)", 1)
    add_body(doc,
        "Ce projet consiste à concevoir et développer HôtelFacile Smart, une application web full-stack "
        "dédiée à la gestion et à la réservation d'hôtels au Maroc. La plateforme met en relation trois "
        "types d'utilisateurs — visiteurs, clients voyageurs, propriétaires d'établissements et "
        "administrateurs — au sein d'un écosystème unique, sécurisé et moderne.")
    add_body(doc,
        "Le problème traité est la fragmentation des outils utilisés par les hôteliers indépendants "
        "et les voyageurs : réservation manuelle, absence de visibilité en ligne, modération des "
        "établissements, suivi des chambres et des avis clients. La solution proposée centralise "
        "l'inscription, la recherche multi-critères, la réservation, la modération administrative, "
        "les notifications en temps réel et la gestion des favoris et des avis.")
    add_body(doc,
        "Les technologies utilisées sont Node.js et Express pour le backend REST, MySQL pour la "
        "persistance des données, HTML/CSS/JavaScript vanilla pour le frontend multi-rôles, JWT et "
        "bcrypt pour la sécurité, Leaflet/OpenStreetMap pour la géolocalisation, et Google OAuth "
        "pour l'authentification sociale.")
    add_body(doc,
        "Les résultats obtenus montrent qu'il est possible de livrer une application cohérente, "
        "responsive et évolutive, avec une charte graphique unifiée (bleu #2F64EA et jaune #FBBF24), "
        "des parcours utilisateurs distincts et une architecture monolithique légère adaptée à un "
        "contexte académique et à un déploiement progressif.")
    add_body(doc,
        "Mots-clés : gestion hôtelière, application web, Node.js, Express.js, MySQL, JWT, OAuth 2.0, "
        "réservation en ligne, tourisme, Maroc, full-stack.")

    add_heading(doc, "Abstract — English", 2)
    add_body(doc,
        "This final year project presents the design and development of HôtelFacile Smart, a full-stack "
        "web application dedicated to hotel management and booking in Morocco. The platform addresses "
        "the fragmentation of tools used by independent hoteliers and travelers by centralizing search, "
        "moderation, reservations, reviews, favorites, and real-time notifications in a single ecosystem.")
    add_body(doc,
        "The solution relies on a Node.js/Express REST backend, a MySQL relational database, and a "
        "multi-portal vanilla HTML/CSS/JavaScript frontend (Public, Client, Owner, Admin). Security is "
        "ensured through JWT authentication, bcrypt password hashing, role-based access control, and "
        "Google OAuth 2.0. Smart features currently rely on heuristics (popularity ranking, fuzzy admin "
        "search, contextual notifications); machine learning-based recommendation is envisioned as a "
        "future extension.")
    add_body(doc,
        "Keywords: hotel management, web application, Node.js, MySQL, JWT, online booking, Morocco.")
    doc.add_page_break()

    # ===================== TABLE DES MATIÈRES =====================
    add_heading(doc, "Table des matières", 1)
    add_body(doc, "Générer automatiquement dans Word : Références → Table des matières → Table automatique.")
    toc_items = [
        "Introduction générale",
        "Chapitre 1 — Contexte et état de l'art",
        "Chapitre 2 — Analyse des besoins",
        "Chapitre 3 — Conception",
        "Chapitre 4 — Développement",
        "Chapitre 5 — Résultats et perspectives",
        "Conclusion générale",
        "Bibliographie",
        "Annexes",
    ]
    for item in toc_items:
        add_bullet(doc, item)
    doc.add_page_break()

    # ===================== INTRODUCTION =====================
    add_heading(doc, "Introduction générale", 1)

    add_heading(doc, "Contexte du projet", 2)
    add_body(doc,
        "Le secteur touristique marocain occupe une place stratégique dans l'économie nationale. "
        "Avec des destinations phares comme Marrakech, Fès, Casablanca, Rabat, Tanger, Agadir, Oujda "
        "et Meknès, la demande de solutions numériques pour découvrir, comparer et réserver des "
        "hébergements ne cesse de croître. Les voyageurs attendent des interfaces claires, rapides "
        "et fiables ; les propriétaires, des outils simples pour publier leurs établissements ; "
        "les administrateurs, un contrôle qualité sur les hôtels mis en ligne.")

    add_heading(doc, "Présentation du domaine", 2)
    add_body(doc,
        "Le domaine visé est celui du tourisme hôtelier et de la digitalisation des services "
        "d'hébergement. Il s'agit de modéliser le cycle de vie d'un hôtel — de sa soumission "
        "par un propriétaire à sa validation par un administrateur, puis à sa consultation "
        "et sa réservation par un client — tout en intégrant des mécanismes de confiance "
        "(modération, avis, notifications).")

    add_heading(doc, "Problématique", 2)
    add_body(doc,
        "Comment concevoir une plateforme web unique capable de : (1) permettre aux voyageurs "
        "de rechercher et réserver des chambres selon plusieurs critères ; (2) offrir aux "
        "propriétaires un espace de gestion complet (hôtels, chambres, photos, réservations, avis) ; "
        "(3) garantir aux administrateurs un contrôle centralisé de la qualité des établissements ; "
        "(4) assurer la sécurité des comptes et la traçabilité des opérations ?")

    add_heading(doc, "Objectifs", 2)
    objectives = [
        "Développer une application web multi-rôles (Public, Client, Propriétaire, Admin).",
        "Implémenter une API REST complète avec authentification JWT et contrôle d'accès par rôle.",
        "Concevoir une base de données relationnelle MySQL normalisée et évolutive.",
        "Proposer une interface utilisateur moderne, cohérente et responsive.",
        "Intégrer des fonctionnalités « intelligentes » : recherche avancée, hôtels populaires, notifications contextuelles.",
        "Documenter l'architecture, les choix techniques et les scénarios de test.",
    ]
    for o in objectives:
        add_bullet(doc, o)

    add_heading(doc, "Structure du rapport", 2)
    add_body(doc,
        "Ce document est organisé en cinq chapitres. Le premier présente le contexte et l'état "
        "de l'art. Le second détaille l'analyse des besoins. Le troisième expose la conception "
        "(architecture, base de données, API). Le quatrième décrit le développement et la sécurité. "
        "Le cinquième analyse les résultats, les tests et les perspectives d'évolution.")
    doc.add_page_break()

    # ===================== CHAPITRE 1 =====================
    add_heading(doc, "Chapitre 1 : Contexte et état de l'art", 1)

    add_heading(doc, "1.1 Présentation du domaine", 2)
    add_body(doc,
        "Le tourisme hôtelier repose sur la mise en relation entre l'offre (établissements, chambres, "
        "tarifs) et la demande (voyageurs, dates, budget). Historiquement, cette relation passait par "
        "des intermédiaires physiques ou des appels téléphoniques. Aujourd'hui, les plateformes web "
        "agrègent l'information, facilitent la comparaison et automatisent la réservation.")
    add_body(doc,
        "Au Maroc, la diversité des types d'hébergement (économique, standard, supérieur, deluxe, "
        "suite, familiale) et la répartition géographique des villes touristiques imposent des "
        "filtres de recherche performants et une cartographie claire. HôtelFacile Smart s'inscrit "
        "dans cette dynamique en ciblant explicitement le contexte marocain.")

    add_heading(doc, "1.2 Problématique", 2)
    add_body(doc, "Description du problème :")
    add_bullet(doc, "Manque d'outils unifiés pour les petits et moyens hôteliers indépendants.")
    add_bullet(doc, "Difficulté pour les voyageurs à comparer prix, capacité et localisation.")
    add_bullet(doc, "Risque de publication d'établissements non conformes sans modération.")
    add_bullet(doc, "Absence de suivi centralisé des réservations, avis et notifications.")
    add_body(doc, "Impact : perte de revenus pour les hôteliers, expérience utilisateur dégradée, faible confiance numérique.")
    add_body(doc, "Importance : une plateforme structurée améliore l'efficacité opérationnelle et la satisfaction client.")

    add_heading(doc, "1.3 Étude des solutions existantes", 2)
    add_body(doc,
        "Avant de concevoir HôtelFacile Smart, nous avons analysé les principales plateformes de "
        "réservation et de gestion hôtelière. Le tableau ci-dessous synthétise leurs forces et limites "
        "au regard de notre cahier des charges.")
    add_table_generic(doc, ["Solution", "Avantages", "Limites"], [
        ("Booking.com", "Catalogue mondial, paiement intégré, notoriété", "Commissions élevées, faible personnalisation locale"),
        ("Expedia", "Offres combinées vol+hôtel, marketing puissant", "Modèle OTA, dépendance aux commissions"),
        ("Airbnb", "UX soignée, avis voyageurs, photos", "Orienté locations, peu adapté à l'hôtellerie classique"),
        ("PMS hôteliers", "Gestion interne complète (réception, facturation)", "Coûteux, complexes, sans vitrine publique"),
        ("Sites statiques d'hôtels", "Autonomie de l'hôtelier", "Pas de recherche transversale ni modération centralisée"),
    ])
    add_body(doc,
        "HôtelFacile Smart se positionne à l'intersection des OTA et des outils de gestion interne : "
        "vitrine publique pour les voyageurs, back-office pour les propriétaires, modération pour "
        "l'administrateur — le tout maîtrisé de bout en bout dans un contexte académique.")

    add_heading(doc, "1.4 Solution proposée", 2)
    add_body(doc,
        "HôtelFacile Smart propose une marketplace hôtelière complète développée en interne. "
        "L'idée principale est de regrouper, dans une seule application, la vitrine publique, "
        "l'espace client, l'espace propriétaire et le back-office administrateur.")
    add_body(doc, "Valeur ajoutée :")
    add_bullet(doc, "Modération admin avant publication des hôtels.")
    add_bullet(doc, "Gestion des chambres avec galerie photos et géolocalisation Leaflet.")
    add_bullet(doc, "Workflow de réservation (en attente → confirmée / refusée / annulée).")
    add_bullet(doc, "Système de favoris, avis clients et notifications en temps réel.")
    add_body(doc,
        "Apport « intelligent » (honnêteté technique) : le terme Smart ne désigne pas aujourd'hui un "
        "module d'apprentissage automatique déployé, mais des mécanismes heuristiques et métier : "
        "classement des hôtels populaires, recherche admin multi-critères avec normalisation des accents, "
        "recherche propriétaire fusionnée (nom d'hôtel, type et capacité de chambre sur une même page), "
        "politique de suppression intelligente (deletionPolicy.js), vérification des chevauchements de "
        "dates de réservation, et notifications contextuelles (notificationService.js). Une évolution "
        "future pourrait intégrer un moteur de recommandation par apprentissage automatique.")

    add_placeholder(doc, "FIGURE 1.1 — Schéma du positionnement de HôtelFacile Smart",
                    "Insérer un diagramme de contexte (acteurs externes, plateforme, base de données).", 2)
    doc.add_page_break()

    # ===================== CHAPITRE 2 =====================
    add_heading(doc, "Chapitre 2 : Analyse des besoins", 1)

    add_heading(doc, "2.1 Identification des acteurs", 2)
    add_body(doc,
        "L'application implique quatre acteurs humains principaux et des systèmes externes, chacun "
        "avec des privilèges distincts.")

    add_heading(doc, "2.1.1 Visiteur (Public)", 3)
    add_body(doc,
        "Internaute non authentifié : page d'accueil, recherche d'hôtels validés par ville/type/budget, "
        "consultation des fiches hôtel et chambre. Pas de réservation ni de favoris sans compte.")

    add_heading(doc, "2.1.2 Client", 3)
    add_body(doc,
        "Voyageur inscrit : recherche avancée, réservation avec dates, historique, modification/annulation "
        "si statut en attente, favoris, avis (note 1–5 + commentaire), notifications sur ses réservations, "
        "connexion email/mot de passe ou Google OAuth.")

    add_heading(doc, "2.1.3 Propriétaire d'hôtel", 3)
    add_body(doc,
        "Hôtelier : création d'hôtels (statut en_attente), gestion chambres et galeries, validation/refus "
        "des réservations, consultation des avis, tableau de bord statistique (revenus, occupation, "
        "répartition par type/capacité). Recherche interne fusionnée sur hotel.html (nom, type, capacité).")

    add_heading(doc, "2.1.4 Administrateur", 3)
    add_body(doc,
        "Superutilisateur : modération des hôtels (valider/refuser), liste filtrée par statut, recherche "
        "avancée multi-critères (nom hôtel, propriétaire, ville) avec normalisation accent-insensible, "
        "statistiques globales de la plateforme.")

    add_heading(doc, "2.1.5 Systèmes externes", 3)
    add_bullet(doc, "Google OAuth 2.0 — authentification sociale vérifiée côté serveur (tokeninfo).")
    add_bullet(doc, "OpenStreetMap / Nominatim — géocodage et carte Leaflet.")
    add_bullet(doc, "Stockage local — uploads/hotels/ et uploads/chambres/.")

    add_placeholder(doc, "FIGURE 2.1 — Diagramme des cas d'utilisation",
                    "Insérer un diagramme UML use case avec les 4 acteurs principaux.", 2)

    add_heading(doc, "2.2 Besoins fonctionnels", 2)

    add_heading(doc, "2.2.1 Module Authentification", 3)
    for f in [
        "Inscription/connexion client et propriétaire (email, bcrypt).",
        "Connexion admin dédiée ; routes séparées par rôle.",
        "Google OAuth 2.0 ; profil, mot de passe, suppression compte (confirmation « SUPPRIMER »).",
    ]:
        add_bullet(doc, f)

    add_heading(doc, "2.2.2 Module Hôtels", 3)
    for f in [
        "CRUD hôtel propriétaire ; workflow en_attente → valide/refuse.",
        "Upload image principale + galerie ; GPS latitude/longitude.",
        "Recherche publique /hotels/search ; hôtels et villes populaires.",
        "Recherche admin /hotels/admin/list ; recherche propriétaire /hotels/searchProprietaire.",
    ]:
        add_bullet(doc, f)

    add_heading(doc, "2.2.3 Module Chambres et Réservations", 3)
    for f in [
        "Types : Économique, Standard, Supérieure, Deluxe, Suite, Familiale.",
        "Réservation avec anti-chevauchement de dates ; statuts en_attente, confirmee, refusee, annulee.",
        "Modification client uniquement si en_attente.",
    ]:
        add_bullet(doc, f)

    add_heading(doc, "2.2.4 Modules Avis, Favoris, Notifications, Statistiques", 3)
    for f in [
        "Avis unique par client/hôtel (ON DUPLICATE KEY UPDATE).",
        "Favoris avec prix minimum et métadonnées.",
        "10+ types de notifications (bienvenue, hotel_valide, reservation_confirmee, avis_nouveau…).",
        "Dashboards propriétaire et admin (agrégations SQL, graphiques).",
    ]:
        add_bullet(doc, f)

    add_heading(doc, "2.3 Besoins non fonctionnels", 2)
    nfunc = [
        ("Performance", "Réponses API rapides, chargement des listes paginées côté client, cache sessionStorage pour l'accueil."),
        ("Sécurité", "JWT, bcrypt, contrôle de rôle, validation uploads (MIME, 5 Mo max)."),
        ("Disponibilité", "Architecture monolithique simple à déployer ; base MySQL relationnelle."),
        ("Ergonomie", "Charte graphique unifiée, navigation cohérente, responsive design."),
        ("Maintenabilité", "Séparation frontend par rôle, routes API modulaires, migrations SQL incrémentales."),
        ("Évolutivité", "Ajout de modules (paiement, IA recommandation) sans refonte totale."),
    ]
    for name, desc in nfunc:
        p = doc.add_paragraph()
        p.add_run(f"{name} : ").bold = True
        p.add_run(desc)

    add_heading(doc, "2.4 User Stories", 2)
    stories = [
        ("US-01", "visiteur", "m'inscrire comme client avec email et mot de passe", "accéder aux réservations"),
        ("US-02", "utilisateur Google", "me connecter en un clic via OAuth", "éviter un nouveau mot de passe"),
        ("US-03", "client", "rechercher à Marrakech par type et budget", "trouver un hébergement adapté"),
        ("US-04", "client", "réserver une chambre avec dates d'arrivée et de départ", "sécuriser mon séjour"),
        ("US-05", "propriétaire", "accepter ou refuser les réservations reçues", "piloter mon occupation"),
        ("US-06", "administrateur", "valider ou refuser les hôtels soumis", "garantir la qualité du catalogue"),
        ("US-07", "client", "noter et commenter un hôtel après séjour", "aider d'autres voyageurs"),
        ("US-08", "propriétaire", "consulter revenus et statistiques de mes chambres", "décider en connaissance de cause"),
        ("US-09", "administrateur", "rechercher un hôtel par nom, propriétaire ou ville", "modérer efficacement le catalogue"),
        ("US-10", "propriétaire", "filtrer mes hôtels par nom, type et capacité depuis hotel.html", "retrouver rapidement une chambre"),
    ]
    for us, role, want, goal in stories:
        add_body(doc, f"{us} — En tant que {role}, je veux {want}, afin de {goal}.")

    add_heading(doc, "2.5 Analyse des données", 2)
    add_body(doc,
        "Les données manipulées sont principalement relationnelles : utilisateurs, hôtels, chambres, "
        "réservations, images, favoris, notifications et avis. Les images sont stockées sur le "
        "disque (dossiers uploads/hotels et uploads/chambres) avec chemins référencés en base.")
    add_bullet(doc, "Type : structuré (SQL) + fichiers binaires (images).")
    add_bullet(doc, "Source : saisie utilisateur, uploads, événements métier (réservations, modération).")
    add_bullet(doc, "Volume : adapté à un prototype académique ; scalable via index MySQL et pagination future.")

    add_placeholder(doc, "TABLEAU 2.1 — Matrice des fonctionnalités par acteur",
                    "Insérer un tableau croisé acteurs × fonctionnalités.", 1)
    doc.add_page_break()

    # ===================== CHAPITRE 3 =====================
    add_heading(doc, "Chapitre 3 : Conception", 1)

    add_heading(doc, "3.1 Architecture globale", 2)
    add_body(doc,
        "L'application suit une architecture three-tier simplifiée :")
    add_bullet(doc, "Couche présentation : frontend HTML/CSS/JS (Public, Client, Propriétaire, Admin + Commun).")
    add_bullet(doc, "Couche métier : API REST Express (backend/server.js, routes modulaires).")
    add_bullet(doc, "Couche données : MySQL (gestion_hotel) + système de fichiers pour les médias.")
    add_body(doc,
        "Le serveur Express sert à la fois l'API (/api) et les fichiers statiques du frontend, "
        "ce qui facilite le développement et la démonstration en local (port 3000).")

    add_placeholder(doc, "FIGURE 3.1 — Architecture 3-tiers de HôtelFacile Smart",
                    "Insérer un schéma : Navigateur → Express → MySQL / Uploads.", 2)

    add_heading(doc, "3.2 Choix technologiques", 2)
    add_placeholder(doc, "LOGOS TECHNO",
                    "Insérer les logos : Node.js, Express, MySQL, JavaScript, HTML5, CSS3, JWT, Leaflet, Google OAuth.", 2)

    add_heading(doc, "3.2.1 Backend", 3)
    add_body(doc,
        "Node.js pour les performances I/O et l'écosystème npm ; Express 5 pour structurer l'API REST "
        "de manière modulaire (un routeur par domaine métier).")

    add_heading(doc, "3.2.2 Base de données", 3)
    add_body(doc,
        "MySQL pour les garanties ACID, les contraintes d'intégrité (FK, UNIQUE, CHECK sur les notes) "
        "et le driver mysql2 avec support des Promises.")

    add_heading(doc, "3.2.3 Frontend", 3)
    add_body(doc,
        "HTML/CSS/JavaScript vanilla : maîtrise complète de la pile, pas de build lourd, portails "
        "Public/Client/Propriétaire/Admin + modules Commun (api.js, notificationsBell.js, hotelMap.js). "
        "Charte visuelle : bleu #2F64EA, jaune #FBBF24, police Plus Jakarta Sans.")

    add_heading(doc, "3.2.4 Authentification et sécurité", 3)
    add_body(doc,
        "JWT stateless (2 h), bcrypt (10 rounds), Google OAuth vérifié serveur, RBAC sur chaque route.")

    add_table_generic(doc, ["Couche", "Technologie", "Rôle"], [
        ("Présentation", "HTML5, CSS3, JS ES6+", "4 portails + Commun"),
        ("API", "Node.js, Express 5", "REST JSON, port 3000"),
        ("Données", "MySQL (gestion_hotel)", "9 tables principales"),
        ("Auth", "JWT, bcrypt, Google OAuth", "Sécurité et rôles"),
        ("Carto", "Leaflet, Nominatim", "Géolocalisation hôtels"),
        ("Médias", "uploads/ local", "Images hôtel/chambre"),
        ("Charte", "#2F64EA / #FBBF24", "Identité HôtelFacile Smart"),
    ], fill="FEF3C7")

    tech = [
        ("Frontend", "HTML5, CSS3, JavaScript vanilla — simplicité, pas de build complexe, contrôle total du DOM."),
        ("Backend", "Node.js + Express 5 — API REST légère, écosystème npm riche."),
        ("Base de données", "MySQL via mysql2 — modèle relationnel adapté aux réservations et relations 1-N."),
        ("Authentification", "JWT (jsonwebtoken) + bcrypt — stateless, standard industrie."),
        ("Cartographie", "Leaflet + Nominatim (OpenStreetMap) — géolocalisation des hôtels."),
        ("OAuth", "Google Identity Services — connexion rapide client/propriétaire."),
        ("Typographie", "Plus Jakarta Sans (Google Fonts) — identité visuelle moderne."),
    ]
    for name, desc in tech:
        p = doc.add_paragraph()
        p.add_run(f"{name} : ").bold = True
        p.add_run(desc)

    add_heading(doc, "3.3 Diagrammes", 2)
    add_body(doc, "Les diagrammes suivants sont à insérer en annexe ou dans ce chapitre :")
    add_bullet(doc, "Diagramme de classes (Utilisateur, Hôtel, Chambre, Réservation, Avis, Notification, Favori).")
    add_bullet(doc, "Diagramme de séquence — Réservation client → validation propriétaire.")
    add_bullet(doc, "Diagramme de séquence — Soumission hôtel → modération admin.")
    add_bullet(doc, "Diagramme de déploiement — Serveur Node, MySQL, dossier uploads.")

    add_placeholder(doc, "FIGURE 3.2 — Diagramme de classes",
                    "Insérer le diagramme UML de classes complet.", 3)
    add_placeholder(doc, "FIGURE 3.3 — Diagramme de séquence (réservation)",
                    "Insérer le diagramme de séquence du workflow de réservation.", 3)

    add_heading(doc, "3.4 Base de données", 2)
    add_body(doc, "Entités principales et relations :")
    tables = [
        "utilisateurs (id, nom, prenom, email, mot_de_passe, role, google_id)",
        "hotels (id, nom, ville, adresse, description, id_proprietaire, statut, image_principale, latitude, longitude)",
        "chambres (id, id_hotel, type, prix, capacite)",
        "reservations (id, id_chambre, id_client, date_debut, date_fin, statut)",
        "hotel_images, chambre_images (galeries)",
        "favoris (id_client, id_hotel — unique)",
        "notifications (type, titre, message, lien, lue, …)",
        "avis (note 1–5, commentaire, id_client, id_hotel — unique)",
    ]
    for t in tables:
        add_bullet(doc, t)
    add_body(doc, "Relations clés : un propriétaire possède plusieurs hôtels ; un hôtel contient plusieurs chambres ; une réservation lie un client à une chambre ; un client peut favoriser plusieurs hôtels et laisser un avis par hôtel.")

    add_placeholder(doc, "FIGURE 3.4 — Schéma relationnel (MLD / MPD)",
                    "Insérer le diagramme entité-association ou le MPD MySQL.", 3)

    add_heading(doc, "3.5 Conception des APIs", 2)
    add_body(doc, "Extrait des endpoints principaux (base : http://localhost:3000/api) :")
    api_rows = [
        ("/auth/client/register", "POST", "Public", "Inscription client"),
        ("/auth/client/login", "POST", "Public", "Connexion client"),
        ("/auth/proprietaire/register", "POST", "Public", "Inscription propriétaire"),
        ("/auth/proprietaire/login", "POST", "Public", "Connexion propriétaire"),
        ("/auth/admin/login", "POST", "Public", "Connexion admin"),
        ("/auth/google", "POST", "Public", "OAuth Google (tokeninfo)"),
        ("/profil", "GET/PUT/DELETE", "Authentifié", "Profil et suppression compte"),
        ("/hotels", "POST", "Propriétaire", "Créer un hôtel (en_attente)"),
        ("/hotels/mes-hotels", "GET", "Propriétaire", "Lister mes hôtels"),
        ("/hotels/mes-hotels/:id", "GET", "Propriétaire", "Détail hôtel propriétaire"),
        ("/hotels/admin/list", "GET", "Admin", "Liste + recherche floue"),
        ("/hotels/admin/:id", "GET", "Admin", "Fiche modération"),
        ("/hotels/:id/valider", "PUT", "Admin", "Valider un hôtel"),
        ("/hotels/:id/refuser", "PUT", "Admin", "Refuser un hôtel"),
        ("/hotels/search", "GET", "Public", "Recherche hôtels + chambres"),
        ("/hotels/searchProprietaire", "GET", "Propriétaire", "Recherche fusionnée"),
        ("/hotels/popular", "GET", "Public", "Hôtels populaires"),
        ("/villes/popular", "GET", "Public", "Villes populaires"),
        ("/hotels/validated", "GET", "Public", "Catalogue validé"),
        ("/chambres", "POST", "Propriétaire", "Ajouter une chambre"),
        ("/chambres/:id_hotel", "GET", "Public/Prop.", "Lister chambres"),
        ("/reservations", "POST", "Client", "Créer réservation"),
        ("/reservations/:id/valider", "PUT", "Propriétaire", "Confirmer"),
        ("/reservations/:id/refuser", "PUT", "Propriétaire", "Refuser"),
        ("/reservations/:id/annuler", "PUT", "Client", "Annuler"),
        ("/reservations/mes-reservations", "GET", "Client", "Historique client"),
        ("/favoris", "GET", "Client", "Liste favoris"),
        ("/favoris/:hotelId", "POST/DELETE", "Client", "Ajouter/retirer favori"),
        ("/notifications", "GET", "Authentifié", "Notifications"),
        ("/notifications/unread-count", "GET", "Authentifié", "Compteur non lues"),
        ("/avis", "POST", "Client", "Publier/mettre à jour avis"),
        ("/hotels/:hotelId/avis", "GET", "Public", "Avis publics d'un hôtel"),
        ("/proprietaire/stats", "GET", "Propriétaire", "Stats agrégées"),
        ("/admin/stats", "GET", "Admin", "Stats globales"),
        ("/hotels/:id/images", "POST/GET", "Mixte", "Galerie images hôtel"),
        ("/chambres/:id/images", "POST/GET", "Mixte", "Galerie images chambre"),
    ]
    add_table_api(doc, api_rows)

    add_heading(doc, "3.6 Conception du module intelligent", 2)
    add_body(doc,
        "Le module « Smart » repose actuellement sur des mécanismes heuristiques et métier plutôt "
        "qu'un modèle d'apprentissage profond :")
    add_bullet(doc, "Classement des hôtels populaires (nombre de réservations confirmées).")
    add_bullet(doc, "Recherche floue admin (normalisation accents, filtres combinés).")
    add_bullet(doc, "Notifications contextuelles déclenchées par événements (notificationService.js).")
    add_bullet(doc, "Recommandations par ville et filtres (type, budget, capacité).")
    add_body(doc,
        "Évolution IA envisagée : moteur de recommandation collaboratif (filtrage basé sur les "
        "réservations et avis), prédiction de taux d'occupation, détection d'anomalies sur les avis.")

    add_placeholder(doc, "FIGURE 3.5 — Flux du module de notifications intelligentes",
                    "Insérer un schéma des hooks notificationService → base → cloche frontend.", 2)
    doc.add_page_break()

    # ===================== CHAPITRE 4 =====================
    add_heading(doc, "Chapitre 4 : Développement", 1)

    add_heading(doc, "4.1 Environnement de développement", 2)
    tools = [
        "IDE : Visual Studio Code / Cursor",
        "Runtime : Node.js",
        "Base : MySQL (gestion_hotel)",
        "Navigateur : Chrome / Edge (DevTools)",
        "Versionnement : Git",
        "API testing : fetch / Thunder Client / Postman",
    ]
    for t in tools:
        add_bullet(doc, t)

    add_heading(doc, "4.2 Implémentation — Structure du projet", 2)
    add_body(doc, "Organisation des dossiers principaux :")
    structure = [
        ("backend/server.js", "Point d'entrée Express, montage des routes et fichiers statiques."),
        ("backend/routes/", "10 routeurs : auth, hotels, chambres, réservations, stats, favoris, notifications, avis, images."),
        ("backend/middleware/authMiddleware.js", "Vérification JWT Bearer."),
        ("backend/utils/", "notificationService, deletionPolicy, hotelImageStorage."),
        ("backend/database/", "Migrations SQL incrémentales (images, favoris, avis, notifications, GPS…)."),
        ("frontend/Public/", "Site vitrine : index, recherche, détail hôtel (lecture seule)."),
        ("frontend/Client/", "Espace client : dashboard, réservation, favoris, avis."),
        ("frontend/Proprietaire/", "Espace propriétaire : hôtels, chambres, réservations, stats."),
        ("frontend/Admin/", "Back-office : modération, recherche, statistiques."),
        ("frontend/Commun/", "Modules partagés : api.js, auth, notifications, carte, galerie, profil."),
        ("frontend/css/style.css", "Feuille principale (~13 000 lignes), charte #2F64EA / #FBBF24."),
    ]
    for path, desc in structure:
        p = doc.add_paragraph()
        p.add_run(f"{path} — ").bold = True
        p.add_run(desc)

    add_heading(doc, "4.2.1 Module Public", 2)
    add_body(doc,
        "Le module Public (frontend/Public/html/index.html) constitue la vitrine de la plateforme. "
        "Il propose un hero immersif, un formulaire de recherche (ville, voyageurs, type, budget) "
        "et des sections destinations / hôtels populaires alimentées par homeData.js.")

    add_heading(doc, "4.2.2 Module Client", 2)
    add_body(doc,
        "L'espace client permet l'inscription, la connexion (email ou Google), la recherche avancée, "
        "la réservation de chambres, la gestion des favoris et des avis, ainsi qu'un tableau de bord "
        "personnalisé avec notifications.")

    add_heading(doc, "4.2.3 Module Propriétaire", 2)
    add_body(doc,
        "Le propriétaire crée ses hôtels via hotelAjouter.html (carte Leaflet, upload photos, chambres initiales). "
        "Après soumission, l'hôtel reste en attente jusqu'à validation admin. Il gère ensuite ses chambres, "
        "valide les réservations et consulte ses statistiques (revenus, répartition par type/capacité).")

    add_heading(doc, "4.2.4 Module Admin", 2)
    add_body(doc,
        "L'administrateur dispose d'un dashboard (statistiques, hôtels récents), d'une page hôtels "
        "avec filtres par statut et recherche avancée, et d'une fiche détail pour valider ou refuser "
        "chaque établissement. La navigation a été unifiée (Accueil, Hôtels, Contact).")

    add_placeholder(doc, "CAPTURES D'ÉCRAN — Interface",
                    "Insérer les captures : accueil public, dashboard client, hotel.html propriétaire, dashboard admin, fiche détail, résultats recherche.", 4)

    add_heading(doc, "4.3 Intégration des fonctionnalités intelligentes", 2)
    add_body(doc,
        "Le service notificationService.js expose une fonction fire() asynchrone et non bloquante, "
        "appelée depuis les routes métier. Types implémentés : bienvenue_client, bienvenue_proprietaire, "
        "hotel_en_attente, hotel_valide, hotel_refuse, reservation_en_attente, reservation_confirmee, "
        "reservation_refusee, reservation_annulee, avis_nouveau. Chaque entrée comporte titre, message "
        "et lien de redirection contextuel.")
    add_body(doc,
        "La recherche admin (/hotels/admin/list) combine filtres statut, nom d'hôtel, prénom/nom "
        "propriétaire et ville avec normalisation accent-insensible (normalizeSearchTerms). "
        "Côté propriétaire, hotel.html regroupe nom d'hôtel, type et capacité dans un bandeau unique "
        "redirigeant vers rechercheResultat.html — ergonomie « recherche fusionnée ».")
    add_body(doc,
        "notificationsBell.js interroge /notifications/unread-count et affiche la liste déroulante ; "
        "homeData.js alimente hôtels/villes populaires ; deletionPolicy.js bloque suppressions si "
        "réservations actives.")
    add_placeholder(doc, "CAPTURE",
                    "Cloche de notifications, résultats recherche admin, bandeau recherche propriétaire fusionné.", 2)

    add_heading(doc, "4.4 Difficultés rencontrées", 2)
    add_body(doc, "Techniques :")
    add_bullet(doc, "Gestion des uploads images en base64 et stockage disque cohérent.")
    add_bullet(doc, "Politique de suppression (deletionPolicy.js) si réservations actives.")
    add_bullet(doc, "Unification CSS multi-rôles sans framework frontend.")
    add_bullet(doc, "Recherche admin floue et filtres combinés côté API + client.")
    add_body(doc, "Organisationnelles :")
    add_bullet(doc, "Répartition des tâches entre modules (Public, Client, Propriétaire, Admin).")
    add_bullet(doc, "Harmonisation UX après itérations (headers, boutons, barres de recherche).")
    add_bullet(doc, "Gestion des migrations SQL sur une base existante.")

    add_heading(doc, "4.5 Sécurité", 2)
    sec = [
        ("Authentification", "JWT signé (2 h), login séparé par rôle, Google OAuth vérifié côté serveur."),
        ("Autorisation", "Middleware verifierToken + contrôles req.user.role dans chaque route sensible."),
        ("Protection des données", "Mots de passe hashés bcrypt (10 rounds), validation MIME/taille uploads."),
        ("Frontend", "Token en localStorage, redirection si rôle incorrect (requireAuth)."),
        ("Suppression compte", "Confirmation « SUPPRIMER », blocage propriétaire si hôtels existants."),
    ]
    for name, desc in sec:
        p = doc.add_paragraph()
        p.add_run(f"{name} : ").bold = True
        p.add_run(desc)

    add_heading(doc, "4.6 Tests effectués", 2)
    add_body(doc, "Tests manuels réalisés sur les parcours critiques :")
    tests = [
        "Inscription et connexion (client, propriétaire, admin).",
        "Création hôtel → modération admin → visibilité publique.",
        "Recherche multi-critères et réservation client.",
        "Validation/refus réservation par le propriétaire.",
        "Favoris, avis, notifications.",
        "Upload/suppression images hôtel et chambre.",
        "Responsive sur mobile et tablette.",
    ]
    for t in tests:
        add_bullet(doc, t)

    add_placeholder(doc, "TABLEAU 4.1 — Plan de tests fonctionnels",
                    "Insérer un tableau : scénario | étapes | résultat attendu | statut.", 2)
    doc.add_page_break()

    # ===================== CHAPITRE 5 =====================
    add_heading(doc, "Chapitre 5 : Résultats et perspectives", 1)

    add_heading(doc, "5.1 Stratégie de test", 2)
    add_body(doc,
        "La stratégie retenue combine tests manuels exploratoires et tests de bout en bout sur "
        "chaque rôle. Des tests unitaires automatisés (Jest/Mocha) pourraient être ajoutés sur "
        "les utilitaires backend (deletionPolicy, notificationService) dans une phase ultérieure.")

    add_heading(doc, "5.2 Résultats", 2)
    add_body(doc, "L'application HôtelFacile Smart atteint les objectifs fixés :")
    results = [
        "Quatre espaces utilisateurs fonctionnels et cohérents visuellement.",
        "API REST documentée implicitement par les routes (~70 endpoints).",
        "Workflow complet de bout en bout : soumission → modération → réservation → avis.",
        "Interface moderne avec charte bleu/jaune unifiée sur tous les rôles.",
        "Notifications et favoris opérationnels.",
    ]
    for r in results:
        add_bullet(doc, r)

    add_heading(doc, "5.2.1 Analyse des résultats", 2)
    add_body(doc,
        "Les retours utilisateurs tests mettent en avant la clarté de la navigation admin et "
        "propriétaire après refonte des headers, la lisibilité des cartes hôtels et l'utilité "
        "des filtres par statut. La recherche fusionnée avec le bandeau d'accueil propriétaire "
        "améliore l'ergonomie.")

    add_heading(doc, "5.2.2 Limites", 2)
    limits = [
        "Pas de paiement en ligne intégré.",
        "Pas de tests automatisés exhaustifs.",
        "Secrets par défaut en développement (JWT_SECRET, credentials MySQL).",
        "Module IA prédictif non implémenté (recommandations heuristiques uniquement).",
        "Pas de déploiement cloud documenté (Docker, CI/CD).",
    ]
    for l in limits:
        add_bullet(doc, l)

    add_heading(doc, "5.2.3 Comparaison avec l'existant", 2)
    add_body(doc,
        "Contrairement aux OTAs internationales, HôtelFacile Smart offre un contrôle total du code, "
        "une modération intégrée et une adaptation au contexte marocain (villes, types de chambres, "
        "UX en français). Le compromis est un catalogue limité au périmètre du projet académique.")

    add_heading(doc, "5.3 Améliorations et extensions futures", 2)
    future = [
        "Intégration Stripe/CMI pour paiement sécurisé.",
        "Module IA de recommandation personnalisée (collaborative filtering).",
        "Application mobile (React Native ou PWA).",
        "Tableau de bord analytics avancé (graphiques, export PDF).",
        "Messagerie client-propriétaire en temps réel.",
        "Internationalisation (arabe, anglais, espagnol).",
        "Déploiement Docker + HTTPS + variables d'environnement production.",
    ]
    for f in future:
        add_bullet(doc, f)

    add_placeholder(doc, "GRAPHIQUES DE RÉSULTATS",
                    "Insérer des graphiques : répartition des statuts hôtels, réservations par mois, notes moyennes.", 2)
    doc.add_page_break()

    # ===================== CONCLUSION =====================
    add_heading(doc, "Conclusion générale", 1)
    add_body(doc,
        "Ce projet de fin d'études a permis de concevoir et de réaliser HôtelFacile Smart, une "
        "plateforme web complète de gestion et de réservation hôtelière adaptée au marché marocain. "
        "En réunissant visiteurs, clients, propriétaires et administrateurs au sein d'une architecture "
        "monolithique claire, nous avons démontré la maîtrise du cycle complet de développement logiciel.")
    add_body(doc, "Apports principaux :")
    add_bullet(doc, "Modélisation UML et base de données relationnelle.")
    add_bullet(doc, "Développement full-stack JavaScript (Node.js + frontend vanilla).")
    add_bullet(doc, "Sécurisation par JWT, bcrypt et contrôle d'accès par rôle.")
    add_bullet(doc, "UX/UI cohérente avec identité visuelle professionnelle.")
    add_body(doc, "Compétences acquises :")
    skills = [
        "Analyse des besoins et rédaction de spécifications.",
        "Conception d'API REST et schéma relationnel.",
        "Développement frontend multi-pages et multi-rôles.",
        "Gestion de projet itérative et résolution de bugs.",
        "Rédaction technique et documentation.",
    ]
    for s in skills:
        add_bullet(doc, s)

    add_heading(doc, "Mot de fin", 2)
    add_body(doc,
        "HôtelFacile Smart démontre qu'une application professionnelle peut être livrée en full-stack "
        "JavaScript sans framework frontend lourd. Les fonctionnalités intelligentes actuelles sont "
        "heuristiques et métier ; l'IA prédictive reste une piste d'évolution. Nous espérons que ce "
        "travail contribuera à la digitalisation des hôteliers indépendants au Maroc.")
    doc.add_page_break()

    # ===================== BIBLIOGRAPHIE =====================
    add_heading(doc, "Bibliographie", 1)
    refs = [
        "Express.js — Documentation officielle : https://expressjs.com/",
        "MySQL — Documentation : https://dev.mysql.com/doc/",
        "Node.js — Documentation : https://nodejs.org/docs/",
        "JWT — RFC 7519 : https://datatracker.ietf.org/doc/html/rfc7519",
        "Leaflet — Bibliothèque cartographique : https://leafletjs.com/",
        "OpenStreetMap / Nominatim — Géocodage : https://nominatim.org/",
        "OWASP — Top 10 Web Application Security Risks : https://owasp.org/",
        "MDN Web Docs — HTML, CSS, JavaScript : https://developer.mozilla.org/",
        "Google Identity — OAuth 2.0 : https://developers.google.com/identity",
        "Booking.com, Airbnb — Analyse comparative des plateformes de réservation (2024–2025).",
        "OWASP REST Security Cheat Sheet : https://cheatsheetseries.owasp.org/cheatsheets/REST_Security_Cheat_Sheet.html",
        "Ministère du Tourisme du Maroc — https://www.tourisme.gov.ma",
        "Jon Duckett — HTML & CSS: Design and Build Websites — Wiley, 2011.",
        "Thomas Connolly & Carolyn Begg — Database Systems — Pearson, 2014.",
    ]
    for i, ref in enumerate(refs, 1):
        add_body(doc, f"[{i}] {ref}")

    doc.add_page_break()

    # ===================== ANNEXES =====================
    add_heading(doc, "Annexes", 1)

    add_heading(doc, "Annexe A — Inventaire des fichiers du projet", 2)
    add_body(doc,
        "Inventaire généré automatiquement (hors node_modules et .git). Les captures et diagrammes "
        "sont à insérer aux emplacements 【 FIGURE 】 et 【 CAPTURE 】 du corps du rapport.")
    project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    files = collect_project_files(project_root)
    # Limiter aux dossiers pertinents pour le rapport
    relevant = [f for f in files if f.startswith(("backend/", "frontend/", "tools/"))]
    chunk_size = 35
    for i in range(0, len(relevant), chunk_size):
        chunk = relevant[i : i + chunk_size]
        add_table_generic(
            doc,
            ["Chemin relatif", "Couche"],
            [(p, "Backend" if p.startswith("backend/") else ("Frontend" if p.startswith("frontend/") else "Outils"))
             for p in chunk],
            fill="F8FAFC",
        )

    add_heading(doc, "Annexe B — Diagrammes UML", 2)
    add_placeholder(doc, "DIAGRAMMES UML COMPLETS",
                    "Insérer : use case, classes, séquences, déploiement, MLD.", 4)

    add_heading(doc, "Annexe C — Captures d'écran", 2)
    add_placeholder(doc, "CAPTURES COMPLÉMENTAIRES",
                    "Insérer toutes les captures par rôle (login, register, détail chambre, notifications, profil…).", 4)

    add_heading(doc, "Annexe D — Extraits de code significatifs", 2)
    add_body(doc, "Exemples de fichiers clés à joindre : server.js, authMiddleware.js, hotelRoutes.js, notificationService.js, api.js.")

    add_heading(doc, "Annexe E — Scripts SQL", 2)
    add_body(doc, "Joindre les fichiers de migration backend/database/*.sql.")

    add_placeholder(doc, "LOGOS TECHNOLOGIES (ANNEXE)",
                    "Planche récapitulative des logos : Node.js, Express, MySQL, JWT, Leaflet, Google, HTML5, CSS3, JavaScript.", 2)

    import shutil
    save_path = OUTPUT_TMP
    try:
        doc.save(OUTPUT)
        save_path = OUTPUT
    except PermissionError:
        doc.save(OUTPUT_TMP)
        try:
            shutil.copy2(OUTPUT_TMP, OUTPUT)
            save_path = OUTPUT
        except PermissionError:
            save_path = OUTPUT_TMP
    print(f"Rapport généré : {save_path}")


if __name__ == "__main__":
    build()
